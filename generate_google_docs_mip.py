from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_google_docs_compatible_mip():
    """Cria um MIP compatível com Google Docs"""
    
    # Criar documento simples
    doc = Document()
    
    # Configurar margens padrão
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título principal
    title = doc.add_heading("MODELO DE INSTRUÇÃO DE PROCEDIMENTO (MIP)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Título do procedimento
    proc_title = doc.add_heading("Procedimento de Configuração de Sistema", 1)
    proc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Visão geral
    doc.add_heading("Visão Geral", 2)
    doc.add_paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
        "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. "
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur."
    )
    
    # Índice
    doc.add_heading("Índice", 2)
    doc.add_paragraph("1. Preparação do Ambiente")
    doc.add_paragraph("2. Configuração Inicial")
    doc.add_paragraph("3. Teste do Sistema")
    doc.add_paragraph("4. Validação Final")
    
    # Quebra de página
    doc.add_page_break()
    
    # Passo 1
    doc.add_heading("1. Preparação do Ambiente", 2)
    doc.add_paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
        "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat."
    )
    
    # Subpassos
    doc.add_paragraph("• Verificar se todos os requisitos estão atendidos")
    doc.add_paragraph("• Confirmar acesso ao sistema")
    doc.add_paragraph("• Preparar credenciais de acesso")
    
    # Passo 2
    doc.add_heading("2. Configuração Inicial", 2)
    doc.add_paragraph(
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. "
        "Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum."
    )
    
    doc.add_paragraph("• Acessar painel de configuração")
    doc.add_paragraph("• Definir parâmetros básicos")
    doc.add_paragraph("• Salvar configurações iniciais")
    
    # Passo 3
    doc.add_heading("3. Teste do Sistema", 2)
    doc.add_paragraph(
        "Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium doloremque laudantium, "
        "totam rem aperiam, eaque ipsa quae ab illo inventore veritatis et quasi architecto beatae vitae dicta sunt explicabo."
    )
    
    doc.add_paragraph("• Executar testes automatizados")
    doc.add_paragraph("• Verificar funcionalidades principais")
    doc.add_paragraph("• Documentar resultados dos testes")
    
    # Passo 4
    doc.add_heading("4. Validação Final", 2)
    doc.add_paragraph(
        "Nemo enim ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit, "
        "sed quia consequuntur magni dolores eos qui ratione voluptatem sequi nesciunt."
    )
    
    doc.add_paragraph("• Revisar todas as configurações")
    doc.add_paragraph("• Confirmar funcionamento correto")
    doc.add_paragraph("• Finalizar procedimento")
    
    # Observações
    doc.add_heading("Observações Importantes", 2)
    doc.add_paragraph("• Semper in massa tempor nec feugiat nisl pretium fusce")
    doc.add_paragraph("• Consectetur adipiscing elit duis tristique sollicitudin")
    doc.add_paragraph("• In hac habitasse platea dictumst vestibulum rhoncus")
    
    # Informações de contato
    doc.add_paragraph()
    contact_para = doc.add_paragraph("Suporte Ramal Virtual - (11) 3090-0900")
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    link_para = doc.add_paragraph("suporte.ramalvirtual.com.br")
    link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_para = doc.add_paragraph("www.ramalvirtual.com.br    CEP:07021-050    Rua Santa Rita de Cássia 155")
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar documento
    output_path = "output/mip_google_docs_compatible.docx"
    doc.save(output_path)
    print(f"MIP compatível com Google Docs gerado: {output_path}")
    print("\nPara usar no Google Docs:")
    print("1. Acesse docs.google.com")
    print("2. Clique em 'Arquivo' > 'Abrir'")
    print("3. Faça upload do arquivo .docx")
    print("4. O Google Docs irá converter automaticamente")

if __name__ == '__main__':
    # Criar diretório output se não existir
    os.makedirs("output", exist_ok=True)
    create_google_docs_compatible_mip() 