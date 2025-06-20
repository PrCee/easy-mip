from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime

def add_hyperlink(paragraph, url, text):
    """Adiciona um hyperlink ao documento"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def create_mip_docx():
    """Cria um MIP em formato DOCX editável"""
    
    # Criar documento
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Definir estilos
    styles = doc.styles
    
    # Estilo para título principal
    title_style = styles.add_style('MIPTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 81, 124)  # Azul corporativo
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(20)
    
    # Estilo para subtítulos
    subtitle_style = styles.add_style('MIPSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Arial'
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.bold = True
    subtitle_style.font.color.rgb = RGBColor(0, 81, 124)
    subtitle_style.paragraph_format.space_after = Pt(12)
    
    # Estilo para passos
    step_style = styles.add_style('MIPStep', WD_STYLE_TYPE.PARAGRAPH)
    step_style.font.name = 'Arial'
    step_style.font.size = Pt(12)
    step_style.paragraph_format.left_indent = Inches(0.2)
    step_style.paragraph_format.space_after = Pt(10)
    
    # Estilo para texto normal
    normal_style = styles.add_style('MIPNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(8)
    
    # Cabeçalho
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "Ramal Virtual - Suporte Técnico"
    header_para.style = subtitle_style
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Título principal
    title = doc.add_paragraph("MODELO DE INSTRUÇÃO DE PROCEDIMENTO (MIP)")
    title.style = title_style
    
    # Título do procedimento
    proc_title = doc.add_paragraph("Procedimento de Configuração de Sistema")
    proc_title.style = title_style
    
    # Visão geral
    doc.add_paragraph("Visão Geral", style='MIPSubtitle')
    overview = doc.add_paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
        "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. "
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur.",
        style='MIPNormal'
    )
    
    # Índice
    doc.add_paragraph("Índice", style='MIPSubtitle')
    doc.add_paragraph("1. Preparação do Ambiente", style='MIPStep')
    doc.add_paragraph("2. Configuração Inicial", style='MIPStep')
    doc.add_paragraph("3. Teste do Sistema", style='MIPStep')
    doc.add_paragraph("4. Validação Final", style='MIPStep')
    
    # Quebra de página
    doc.add_page_break()
    
    # Passo 1
    doc.add_paragraph("1. Preparação do Ambiente", style='MIPSubtitle')
    step1_desc = doc.add_paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
        "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.",
        style='MIPNormal'
    )
    
    # Subpassos
    doc.add_paragraph("• Verificar se todos os requisitos estão atendidos", style='MIPStep')
    doc.add_paragraph("• Confirmar acesso ao sistema", style='MIPStep')
    doc.add_paragraph("• Preparar credenciais de acesso", style='MIPStep')
    
    # Passo 2
    doc.add_paragraph("2. Configuração Inicial", style='MIPSubtitle')
    step2_desc = doc.add_paragraph(
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. "
        "Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.",
        style='MIPNormal'
    )
    
    doc.add_paragraph("• Acessar painel de configuração", style='MIPStep')
    doc.add_paragraph("• Definir parâmetros básicos", style='MIPStep')
    doc.add_paragraph("• Salvar configurações iniciais", style='MIPStep')
    
    # Passo 3
    doc.add_paragraph("3. Teste do Sistema", style='MIPSubtitle')
    step3_desc = doc.add_paragraph(
        "Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium doloremque laudantium, "
        "totam rem aperiam, eaque ipsa quae ab illo inventore veritatis et quasi architecto beatae vitae dicta sunt explicabo.",
        style='MIPNormal'
    )
    
    doc.add_paragraph("• Executar testes automatizados", style='MIPStep')
    doc.add_paragraph("• Verificar funcionalidades principais", style='MIPStep')
    doc.add_paragraph("• Documentar resultados dos testes", style='MIPStep')
    
    # Passo 4
    doc.add_paragraph("4. Validação Final", style='MIPSubtitle')
    step4_desc = doc.add_paragraph(
        "Nemo enim ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit, "
        "sed quia consequuntur magni dolores eos qui ratione voluptatem sequi nesciunt.",
        style='MIPNormal'
    )
    
    doc.add_paragraph("• Revisar todas as configurações", style='MIPStep')
    doc.add_paragraph("• Confirmar funcionamento correto", style='MIPStep')
    doc.add_paragraph("• Finalizar procedimento", style='MIPStep')
    
    # Observações
    doc.add_paragraph("Observações Importantes", style='MIPSubtitle')
    doc.add_paragraph("• Semper in massa tempor nec feugiat nisl pretium fusce", style='MIPStep')
    doc.add_paragraph("• Consectetur adipiscing elit duis tristique sollicitudin", style='MIPStep')
    doc.add_paragraph("• In hac habitasse platea dictumst vestibulum rhoncus", style='MIPStep')
    
    # Rodapé com informações de contato
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Suporte Ramal Virtual - (11) 3090-0900"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar link para suporte online
    link_para = footer.add_paragraph()
    add_hyperlink(link_para, "https://suporte.ramalvirtual.com.br", "suporte.ramalvirtual.com.br")
    link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações da empresa
    company_para = footer.add_paragraph("www.ramalvirtual.com.br    CEP:07021-050    Rua Santa Rita de Cássia 155")
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar documento
    output_path = "output/lorem_mip_example.docx"
    doc.save(output_path)
    print(f"MIP editável gerado com sucesso: {output_path}")
    print("Este arquivo pode ser aberto e editado no:")
    print("- Microsoft Word Online")
    print("- Google Docs")
    print("- Microsoft Word Desktop")

if __name__ == '__main__':
    # Criar diretório output se não existir
    os.makedirs("output", exist_ok=True)
    create_mip_docx() 